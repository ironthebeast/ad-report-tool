"""
신고서 생성 모듈
- 공정거래위원회 부당 표시·광고 신고서 양식을 DOCX로 생성
- 양식 구조: 신고인 정보, 피신고인 정보, 위반행위 내용, 증빙자료
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime
import os


def set_cell_border(cell, **kwargs):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        from lxml import etree
        tcBorders = etree.SubElement(tcPr, qn('w:tcBorders'))
    for edge, val in kwargs.items():
        element = tcBorders.find(qn(f'w:{edge}'))
        if element is None:
            from lxml import etree
            element = etree.SubElement(tcBorders, qn(f'w:{edge}'))
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        element.set(qn('w:space'), val.get('space', '0'))


def generate_report(data: dict, save_path: str) -> str:
    """
    공정위 부당 표시·광고 신고서를 DOCX로 생성

    data 구조:
    {
        'reporter': {
            'name': '홍길동',
            'birth_date': '1990-01-01',
            'address': '서울시 강남구...',
            'phone': '010-1234-5678',
            'email': 'hong@email.com',
        },
        'respondent': {
            'business_name': '○○컴퍼니',
            'representative': '김○○',
            'address': '서울시 서초구...',
            'phone': '02-1234-5678',
            'website': 'https://...',
        },
        'violation': {
            'type': '경제적 이해관계 미표시',
            'media': 'SNS (인스타그램)',
            'date': '2026-02-01',
            'url': 'https://instagram.com/p/...',
            'description': '어필리에이트 링크가...',
            'legal_basis': '표시·광고의 공정화에 관한 법률 제3조',
        },
        'evidence': {
            'screenshot_path': '/path/to/screenshot.png',
            'analysis_text': '자동 분석 결과...',
            'affiliate_indicators': ['어필리에이트 링크 3개 발견'],
            'additional_notes': '추가 참고사항',
        }
    }
    """
    doc = Document()

    # ── 페이지 설정 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # ── 스타일 설정 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ══════════════════════════════════════
    # 제목
    # ══════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(4)
    run = title.add_run('부당한 표시·광고 신고서')
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # 법적 근거
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.space_after = Pt(16)
    run = sub.add_run('「표시·광고의 공정화에 관한 법률」 제16조 및 같은 법 시행령 제12조')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ══════════════════════════════════════
    # 1. 신고인 정보
    # ══════════════════════════════════════
    _section_header(doc, '1. 신고인 정보')
    reporter = data.get('reporter', {})
    table1 = doc.add_table(rows=5, cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(table1)
    _fill_row(table1, 0, '성명 *', reporter.get('name', ''))
    _fill_row(table1, 1, '생년월일', reporter.get('birth_date', ''))
    _fill_row(table1, 2, '주소 *', reporter.get('address', ''))
    _fill_row(table1, 3, '전화번호 *', reporter.get('phone', ''))
    _fill_row(table1, 4, '이메일', reporter.get('email', ''))
    # 열 너비
    for row in table1.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12.5)
    doc.add_paragraph().space_after = Pt(8)

    # ══════════════════════════════════════
    # 2. 피신고인 정보
    # ══════════════════════════════════════
    _section_header(doc, '2. 피신고인(사업자) 정보')
    respondent = data.get('respondent', {})
    table2 = doc.add_table(rows=5, cols=2)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(table2)
    _fill_row(table2, 0, '사업자명/계정명 *', respondent.get('business_name', ''))
    _fill_row(table2, 1, '대표자/운영자', respondent.get('representative', ''))
    _fill_row(table2, 2, '주소/소재지', respondent.get('address', ''))
    _fill_row(table2, 3, '전화번호', respondent.get('phone', ''))
    _fill_row(table2, 4, '웹사이트/SNS', respondent.get('website', ''))
    for row in table2.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12.5)
    doc.add_paragraph().space_after = Pt(8)

    # ══════════════════════════════════════
    # 3. 위반행위 내용
    # ══════════════════════════════════════
    _section_header(doc, '3. 위반행위의 내용')
    violation = data.get('violation', {})
    table3 = doc.add_table(rows=6, cols=2)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    _style_table(table3)
    _fill_row(table3, 0, '위반 유형 *', violation.get('type', ''))
    _fill_row(table3, 1, '광고 매체 *', violation.get('media', ''))
    _fill_row(table3, 2, '광고 일자 *', violation.get('date', ''))
    _fill_row(table3, 3, '광고 URL *', violation.get('url', ''))
    _fill_row(table3, 4, '관련 법률', violation.get('legal_basis', ''))
    for row in table3.rows:
        row.cells[0].width = Cm(3.5)
        row.cells[1].width = Cm(12.5)

    # 위반 상세 설명 (큰 셀)
    desc_cell = table3.cell(5, 0)
    _set_cell_text(desc_cell, '위반행위\n상세 설명 *', bold=True, bg='F2F2F2')
    desc_val = table3.cell(5, 1)
    _set_cell_text(desc_val, violation.get('description', ''))
    desc_cell.width = Cm(3.5)
    desc_val.width = Cm(12.5)
    # 높이 확보
    for p in desc_val.paragraphs:
        p.paragraph_format.space_after = Pt(60)

    doc.add_paragraph().space_after = Pt(8)

    # ══════════════════════════════════════
    # 4. 증빙자료
    # ══════════════════════════════════════
    _section_header(doc, '4. 증빙자료')
    evidence = data.get('evidence', {})

    # 자동 분석 결과
    if evidence.get('analysis_text'):
        p = doc.add_paragraph()
        run = p.add_run('[ 자동 분석 결과 ]')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        p = doc.add_paragraph(evidence['analysis_text'])
        p.style.font.size = Pt(9)
        p.space_after = Pt(8)

    # 어필리에이트 지표
    indicators = evidence.get('affiliate_indicators', [])
    if indicators:
        p = doc.add_paragraph()
        run = p.add_run('[ 탐지된 어필리에이트 지표 ]')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        for ind in indicators:
            doc.add_paragraph(f'  • {ind}', style='List Bullet')

    # 스크린샷 첨부
    screenshot = evidence.get('screenshot_path')
    if screenshot and os.path.exists(screenshot):
        p = doc.add_paragraph()
        p.space_before = Pt(12)
        run = p.add_run('[ 화면 캡처 증거 ]')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cap_run = cap_p.add_run(f'캡처 URL: {evidence.get("url", "")}\n캡처 일시: {evidence.get("captured_at", "")}')
        cap_run.font.size = Pt(8)
        cap_run.font.color.rgb = RGBColor(100, 100, 100)

        try:
            doc.add_picture(screenshot, width=Inches(5.5))
        except Exception:
            doc.add_paragraph('[스크린샷 파일을 첨부할 수 없습니다. 별도 파일로 제출해주세요.]')

    # 추가 스크린샷
    extra_shots = evidence.get('extra_screenshots', [])
    for idx, shot_path in enumerate(extra_shots):
        if shot_path and os.path.exists(shot_path):
            p = doc.add_paragraph()
            p.space_before = Pt(8)
            run = p.add_run(f'[ 추가 증거 스크린샷 {idx + 2} ]')
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = '맑은 고딕'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            try:
                doc.add_picture(shot_path, width=Inches(5.5))
            except Exception:
                doc.add_paragraph(f'[스크린샷 {idx + 2} 첨부 실패. 별도 파일로 제출해주세요.]')

    # 추가 참고사항
    if evidence.get('additional_notes'):
        p = doc.add_paragraph()
        p.space_before = Pt(12)
        run = p.add_run('[ 추가 참고사항 ]')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        doc.add_paragraph(evidence['additional_notes'])

    # ══════════════════════════════════════
    # 5. 제출 안내
    # ══════════════════════════════════════
    doc.add_page_break()
    _section_header(doc, '[ 신고서 제출 안내 ]')

    guide_items = [
        ('제출 방법 1 — 국민신문고 (온라인)', '1. epeople.go.kr 접속 → 로그인\n2. 민원신청 → 일반민원\n3. 기관: "공정거래위원회" 선택\n4. 이 신고서(DOCX/PDF) + 증거 스크린샷 첨부\n5. 제출'),
        ('제출 방법 2 — 우편', '피신고업체 관할 지방공정거래사무소로 발송\n• 서울사무소: 서울 중구 세종대로 39\n• 부산사무소: 부산 연제구 중앙대로 1001\n• 대전사무소: 대전 서구 청사로 189\n• 광주사무소: 광주 서구 내방로 111\n• 대구사무소: 대구 수성구 동대구로 340'),
        ('제출 방법 3 — 전화 상담', '공정거래위원회 민원상담: 1670-0007\n소비자 상담: 1372'),
        ('참고사항', '• 신고 후 조사 개시까지 1~3개월 소요될 수 있습니다.\n• 증빙자료는 많을수록 좋습니다.\n• 위반 콘텐츠가 삭제될 수 있으니 스크린샷을 반드시 보관하세요.\n• 이 신고서는 한글(HWP) 또는 Word에서 열어 수정할 수 있습니다.'),
    ]

    for title, content in guide_items:
        p = doc.add_paragraph()
        p.space_before = Pt(8)
        run = p.add_run(f'▸ {title}')
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        p2 = doc.add_paragraph(content)
        p2.paragraph_format.left_indent = Cm(0.5)
        for run2 in p2.runs:
            run2.font.size = Pt(9)
            run2.font.name = '맑은 고딕'

    # ── 날짜 & 서명 ──
    doc.add_paragraph().space_after = Pt(20)
    today = datetime.now().strftime('%Y년 %m월 %d일')
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_p.add_run(today)
    run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_p.space_before = Pt(16)
    run = sign_p.add_run(f'신고인: {reporter.get("name", "___________")}  (서명 또는 인)')
    run.font.size = Pt(11)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    dest_p = doc.add_paragraph()
    dest_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dest_p.space_before = Pt(24)
    run = dest_p.add_run('공정거래위원회 위원장 귀하')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ── 저장 ──
    doc.save(save_path)
    return save_path


# ─── Helper 함수들 ───

def _section_header(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    run.font.color.rgb = RGBColor(26, 26, 46)


def _style_table(table):
    table.style = 'Table Grid'
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)


def _fill_row(table, row_idx, label, value):
    label_cell = table.cell(row_idx, 0)
    _set_cell_text(label_cell, label, bold=True, bg='F2F2F2')
    value_cell = table.cell(row_idx, 1)
    _set_cell_text(value_cell, value)


def _set_cell_text(cell, text, bold=False, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.bold = bold
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if bg:
        shading = cell._element.get_or_add_tcPr()
        from lxml import etree
        shd = etree.SubElement(shading, qn('w:shd'))
        shd.set(qn('w:fill'), bg)
        shd.set(qn('w:val'), 'clear')
